from docx import Document
from docx.shared import Pt

# Create a new document to fill the report content
report = Document()

def add_heading(text, level=1):
    report.add_heading(text, level=level)

def add_paragraph(text):
    para = report.add_paragraph(text)
    para.style.font.size = Pt(12)

# Title page placeholders
add_heading("Lebanese University", level=1)
add_heading("Faculty of Information, Branch I", level=2)
report.add_paragraph("\nUNDERGRADUATE INTERNSHIP REPORT\nBachelor of Data Science")
report.add_paragraph("\nStudent name: ___________________")
report.add_paragraph("Student ID: _____________________")
report.add_paragraph("Academic year: __________________")
report.add_paragraph("Internship period: ______________")
report.add_paragraph("Company name: ___________________")
report.add_paragraph("Institution Supervisor: __________")
report.add_paragraph("Academic Supervisor: ____________")

# Acknowledgment
add_heading("Acknowledgment")
add_paragraph(
    "I would like to express my deepest gratitude to my academic supervisor and my internship mentors "
    "for their continuous guidance and support throughout my internship. I also thank the organization for "
    "providing me with this valuable opportunity to explore real-world applications of data science and business intelligence."
)

# Executive Summary
add_heading("Executive Summary")
add_paragraph(
    "During my internship, I worked in a role focused on data analysis and dashboard design using Power BI. "
    "The goal was to help the organization monitor key metrics and improve internal decision-making. In addition to Power BI work, "
    "I collaborated on automating data ingestion via APIs and built a prototype login system based on QR code scanning."
)
add_paragraph(
    "My main contributions included creating interactive dashboards, automating data cleaning steps, integrating external APIs, and building a mock scanning app "
    "that simulated user authentication via QR codes. These solutions improved internal efficiency and data visibility for the company."
)
add_paragraph(
    "The company benefited from having clearer KPIs, automated reporting processes, and a prototype that could evolve into a secure authentication system."
)

# Introduction
add_heading("Introduction")
add_paragraph(
    "This report details my internship experience where I was assigned a variety of tasks centered on data visualization and automation. "
    "The primary objective was to contribute actionable insights using Power BI dashboards and explore enhancements using automation tools and APIs."
)
add_paragraph(
    "I followed an agile workflow in close collaboration with my supervisor. The internship included weekly meetings, requirement gathering, iterative development, "
    "testing, and documentation of solutions."
)
add_paragraph(
    "Beyond data visualization, I worked with APIs to automate data imports and created a prototype web app that allows login through scanning, showcasing how data-driven "
    "authentication might be implemented. These contributions supported company decisions and opened new opportunities for technical solutions."
)

# Internship Description
add_heading("Internship Description")
add_heading("Company/Institution", level=2)
add_paragraph(
    "The company where I conducted my internship is involved in analytics and digital solutions for internal process optimization. "
    "While it was not a large organization, the environment was dynamic and innovation-driven."
)

add_heading("Methodology", level=2)
add_paragraph(
    "The internship followed a task-based approach. Each task was structured with a defined goal, tools to use, and expected outcomes. "
    "For Power BI, I worked with Excel, SQL views, and dataflows. For automation, I explored API integration using Python scripts and UiPath. "
    "The QR scan login prototype was built using Streamlit and Firebase as backend support."
)

add_heading("Activities Performed", level=2)
add_paragraph(
    "- Developed a dynamic Power BI dashboard to track KPIs across departments.\n"
    "- Cleaned and transformed Excel datasets, and automated repetitive steps in Power Query.\n"
    "- Wrote DAX measures to handle YoY growth, request timing, and outlier detection.\n"
    "- Connected to APIs to fetch data automatically and update dashboards.\n"
    "- Built a prototype login system where users could scan a QR code to authenticate, simulating session-based security."
)

add_heading("Results", level=2)
add_paragraph(
    "The dashboards created were used to present insights to decision-makers, offering an organized view of metrics such as request frequency, floor-wise performance, and command types. "
    "The scanning login app served as a proof-of-concept for potential future implementations in authentication."
)

add_heading("Challenges", level=2)
add_paragraph(
    "One challenge was handling different data formats and ensuring all transformations were applied consistently across multiple Excel files. "
    "In building the mock app, I also had to understand session security and implement TOTP-like logic to simulate secure login flows."
)

add_heading("Final Status of the Project", level=2)
add_paragraph(
    "The dashboards were completed and handed over for daily operational use. The mock login system was finalized as a prototype. "
    "Future work may involve turning this prototype into a production-ready authentication system and expanding dashboard coverage."
)

# Reflection
add_heading("Reflection")
add_heading("Experience Gained", level=2)
add_paragraph(
    "This internship allowed me to deepen my understanding of data pipelines, dashboards, and APIs. I became more confident in using Power BI, DAX, and Streamlit, and I learned how to manage tasks with real business value."
)

add_heading("Skills and Responsibilities", level=2)
add_paragraph(
    "I was trusted to handle live data, design visual interfaces, and propose new features. My supervisors gave me autonomy to research and test ideas, especially in the QR login prototype."
)

add_heading("Professional Understanding", level=2)
add_paragraph(
    "I learned how to balance technical execution with business understanding, how to present findings effectively, and how to document work clearly."
)

add_heading("Future Impact", level=2)
add_paragraph(
    "This experience has reinforced my interest in business intelligence and full-stack development. It helped me see the importance of APIs and data automation in shaping modern software."
)

add_heading("Internship Objectives vs. Achievements", level=2)
add_paragraph(
    "The internship exceeded my expectations. While I expected to work mainly on dashboards, I also got to build and test software, giving me a broader skillset and insight into digital transformation."
)

# Save the completed report
output_path = "DS_Internship_Report.docx"
report.save(output_path)
output_path
